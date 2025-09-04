"""Enhanced Word extractor supporting both DOC and DOCX files."""

import datetime
import logging
import mimetypes
import os
import re
import subprocess
import tempfile
import uuid
from urllib.parse import urlparse
from xml.etree import ElementTree

import requests
from docx import Document as DocxDocument

from configs import dify_config
from core.helper import ssrf_proxy
from core.rag.extractor.extractor_base import BaseExtractor
from core.rag.models.document import Document
from extensions.ext_database import db
from extensions.ext_storage import storage
from models.enums import CreatedByRole
from models.model import UploadFile

logger = logging.getLogger(__name__)


class EnhancedWordExtractor(BaseExtractor):
    """Load both doc and docx files.

    Args:
        file_path: Path to the file to load.
    """

    def __init__(self, file_path: str, tenant_id: str, user_id: str):
        """Initialize with file path."""
        self.file_path = file_path
        self.tenant_id = tenant_id
        self.user_id = user_id

        if "~" in self.file_path:
            self.file_path = os.path.expanduser(self.file_path)

        # If the file is a web path, download it to a temporary file, and use that
        if not os.path.isfile(self.file_path) and self._is_valid_url(self.file_path):
            r = requests.get(self.file_path)

            if r.status_code != 200:
                raise ValueError(f"Check the url of your file; returned status code {r.status_code}")

            self.web_path = self.file_path
            self.temp_file = tempfile.NamedTemporaryFile()
            self.temp_file.write(r.content)
            self.file_path = self.temp_file.name
        elif not os.path.isfile(self.file_path):
            raise ValueError(f"File path {self.file_path} is not a valid file or url")

    def __del__(self) -> None:
        if hasattr(self, "temp_file"):
            self.temp_file.close()

    def extract(self) -> list[Document]:
        """Load given path as documents."""
        _, extension = os.path.splitext(str(self.file_path))
        
        if extension.lower() == '.doc':
            # 处理.doc文件
            content = self._parse_doc(self.file_path)
        elif extension.lower() == '.docx':
            # 处理.docx文件
            content = self.parse_docx(self.file_path, "storage")
        else:
            raise ValueError(f"Unsupported file type: {extension}")
        
        return [
            Document(
                page_content=content,
                metadata={"source": self.file_path},
            )
        ]

    def _parse_doc(self, doc_path: str) -> str:
        """Parse .doc file using various local methods."""
        try:
            logger.info("Using LibreOffice to convert .doc file")
            
            with tempfile.TemporaryDirectory() as temp_dir:
                # 转换为docx
                subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'docx',
                    '--outdir', temp_dir, doc_path
                ], check=True, capture_output=True)
                
                # 获取转换后的文件
                base_name = os.path.splitext(os.path.basename(doc_path))[0]
                docx_path = os.path.join(temp_dir, f"{base_name}.docx")
                
                if os.path.exists(docx_path):
                    # 使用现有的docx解析方法
                    return self.parse_docx(docx_path, "storage")
                else:
                    raise FileNotFoundError("Converted docx file not found")
        except ImportError:
            logger.warning("textract not available, trying alternative methods")
        except Exception as e:
            logger.warning(f"textract failed: {e}, trying alternative methods")
            raise ValueError("无法解析.doc文件。")

    @staticmethod
    def _is_valid_url(url: str) -> bool:
        """Check if the url is valid."""
        parsed = urlparse(url)
        return bool(parsed.netloc) and bool(parsed.scheme)

    def _extract_images_from_docx(self, doc, image_folder):
        """Extract images from docx file (only works for docx)."""
        os.makedirs(image_folder, exist_ok=True)
        image_count = 0
        image_map = {}

        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
                if rel.is_external:
                    url = rel.target_ref
                    response = ssrf_proxy.get(url)
                    if response.status_code == 200:
                        image_ext = mimetypes.guess_extension(response.headers["Content-Type"])
                        if image_ext is None:
                            continue
                        file_uuid = str(uuid.uuid4())
                        file_key = "image_files/" + self.tenant_id + "/" + file_uuid + "." + image_ext
                        mime_type, _ = mimetypes.guess_type(file_key)
                        storage.save(file_key, response.content)
                    else:
                        continue
                else:
                    image_ext = rel.target_ref.split(".")[-1]
                    if image_ext is None:
                        continue
                    file_uuid = str(uuid.uuid4())
                    file_key = "image_files/" + self.tenant_id + "/" + file_uuid + "." + image_ext
                    mime_type, _ = mimetypes.guess_type(file_key)

                    storage.save(file_key, rel.target_part.blob)
                
                # save file to db
                upload_file = UploadFile(
                    tenant_id=self.tenant_id,
                    storage_type=dify_config.STORAGE_TYPE,
                    key=file_key,
                    name=file_key,
                    size=0,
                    extension=str(image_ext),
                    mime_type=mime_type or "",
                    created_by=self.user_id,
                    created_by_role=CreatedByRole.ACCOUNT,
                    created_at=datetime.datetime.now(datetime.UTC).replace(tzinfo=None),
                    used=True,
                    used_by=self.user_id,
                    used_at=datetime.datetime.now(datetime.UTC).replace(tzinfo=None),
                )

                db.session.add(upload_file)
                db.session.commit()
                image_map[rel.target_part] = (
                    f"![image]({dify_config.CONSOLE_API_URL}/files/{upload_file.id}/file-preview)"
                )

        return image_map

    def _table_to_markdown(self, table, image_map):
        """Convert table to markdown format."""
        markdown = []
        total_cols = max(len(row.cells) for row in table.rows)

        header_row = table.rows[0]
        headers = self._parse_row(header_row, image_map, total_cols)
        markdown.append("| " + " | ".join(headers) + " |")
        markdown.append("| " + " | ".join(["---"] * total_cols) + " |")

        for row in table.rows[1:]:
            row_cells = self._parse_row(row, image_map, total_cols)
            markdown.append("| " + " | ".join(row_cells) + " |")
        return "\n".join(markdown)

    def _parse_row(self, row, image_map, total_cols):
        """Parse table row."""
        row_cells = [""] * total_cols
        col_index = 0
        for cell in row.cells:
            while col_index < total_cols and row_cells[col_index] != "":
                col_index += 1
            if col_index >= total_cols:
                break
            cell_content = self._parse_cell(cell, image_map).strip()
            cell_colspan = cell.grid_span or 1
            for i in range(cell_colspan):
                if col_index + i < total_cols:
                    row_cells[col_index + i] = cell_content if i == 0 else ""
            col_index += cell_colspan
        return row_cells

    def _parse_cell(self, cell, image_map):
        """Parse table cell."""
        cell_content = []
        for paragraph in cell.paragraphs:
            parsed_paragraph = self._parse_cell_paragraph(paragraph, image_map)
            if parsed_paragraph:
                cell_content.append(parsed_paragraph)
        unique_content = list(dict.fromkeys(cell_content))
        return " ".join(unique_content)

    def _parse_cell_paragraph(self, paragraph, image_map):
        """Parse cell paragraph."""
        paragraph_content = []
        for run in paragraph.runs:
            if run.element.xpath(".//a:blip"):
                for blip in run.element.xpath(".//a:blip"):
                    image_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if not image_id:
                        continue
                    image_part = paragraph.part.rels[image_id].target_part

                    if image_part in image_map:
                        image_link = image_map[image_part]
                        paragraph_content.append(image_link)
            else:
                paragraph_content.append(run.text)
        return "".join(paragraph_content).strip()

    def _parse_paragraph(self, paragraph, image_map):
        """Parse paragraph."""
        paragraph_content = []
        for run in paragraph.runs:
            if run.element.xpath(".//a:blip"):
                for blip in run.element.xpath(".//a:blip"):
                    embed_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if embed_id:
                        rel_target = run.part.rels[embed_id].target_ref
                        if rel_target in image_map:
                            paragraph_content.append(image_map[rel_target])
            if run.text.strip():
                paragraph_content.append(run.text.strip())
        return " ".join(paragraph_content) if paragraph_content else ""

    def parse_docx(self, docx_path, image_folder):
        """Parse docx file with full formatting support."""
        doc = DocxDocument(docx_path)
        os.makedirs(image_folder, exist_ok=True)

        content = []
        image_map = self._extract_images_from_docx(doc, image_folder)

        hyperlinks_url = None
        url_pattern = re.compile(r"http://[^\s+]+//|https://[^\s+]+")
        
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text and hyperlinks_url:
                    result = f"  [{run.text}]({hyperlinks_url})  "
                    run.text = result
                    hyperlinks_url = None
                if "HYPERLINK" in run.element.xml:
                    try:
                        xml = ElementTree.XML(run.element.xml)
                        x_child = [c for c in xml.iter() if c is not None]
                        for x in x_child:
                            if x_child is None:
                                continue
                            if x.tag.endswith("instrText"):
                                if x.text is None:
                                    continue
                                for i in url_pattern.findall(x.text):
                                    hyperlinks_url = str(i)
                    except Exception:
                        logger.exception("Failed to parse HYPERLINK xml")

        def parse_paragraph(paragraph):
            paragraph_content = []
            for run in paragraph.runs:
                if hasattr(run.element, "tag") and isinstance(run.element.tag, str) and run.element.tag.endswith("r"):
                    drawing_elements = run.element.findall(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                    )
                    for drawing in drawing_elements:
                        blip_elements = drawing.findall(
                            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                        )
                        for blip in blip_elements:
                            embed_id = blip.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                            )
                            if embed_id:
                                image_part = doc.part.related_parts.get(embed_id)
                                if image_part in image_map:
                                    paragraph_content.append(image_map[image_part])
                if run.text.strip():
                    paragraph_content.append(run.text.strip())
            return "".join(paragraph_content) if paragraph_content else ""

        paragraphs = doc.paragraphs.copy()
        tables = doc.tables.copy()
        for element in doc.element.body:
            if hasattr(element, "tag"):
                if isinstance(element.tag, str) and element.tag.endswith("p"):  # paragraph
                    para = paragraphs.pop(0)
                    parsed_paragraph = parse_paragraph(para)
                    if parsed_paragraph.strip():
                        content.append(parsed_paragraph)
                    else:
                        content.append("\n")
                elif isinstance(element.tag, str) and element.tag.endswith("tbl"):  # table
                    table = tables.pop(0)
                    content.append(self._table_to_markdown(table, image_map))
        return "\n".join(content)

    def _convert_doc_to_docx_and_parse(self, doc_path: str) -> str:
        """Convert .doc to .docx using LibreOffice and then parse."""
        import subprocess
        import tempfile
        
        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                # 使用LibreOffice转换
                subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'docx',
                    '--outdir', temp_dir, doc_path
                ], check=True, capture_output=True)
                
                # 获取转换后的文件
                base_name = os.path.splitext(os.path.basename(doc_path))[0]
                docx_path = os.path.join(temp_dir, f"{base_name}.docx")
                
                if os.path.exists(docx_path):
                    # 使用现有的docx解析方法
                    return self.parse_docx(docx_path, "storage")
                else:
                    raise FileNotFoundError("Converted docx file not found")
                    
            except subprocess.CalledProcessError as e:
                logger.error(f"LibreOffice conversion failed: {e}")
                raise
            except Exception as e:
                logger.error(f"Conversion process failed: {e}")
                raise

    def _simple_doc_text_extract(self, doc_path: str) -> str:
        """Simple text extraction from .doc using binary parsing (fallback)."""
        try:
            with open(doc_path, 'rb') as f:
                content = f.read()
                # 简单的二进制文本提取（不完美但可用）
                text = content.decode('utf-8', errors='ignore')
                # 清理非文本字符
                import string
                printable = set(string.printable)
                cleaned_text = ''.join(filter(lambda x: x in printable, text))
                # 移除过多的空白字符
                cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
                return cleaned_text.strip()
        except Exception as e:
            logger.error(f"Simple binary extraction failed: {e}")
            raise ValueError("无法解析.doc文件，请转换为.docx格式或安装相应的解析工具")